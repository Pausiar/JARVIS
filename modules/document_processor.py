"""
J.A.R.V.I.S. — Document Processor Module
Procesamiento de PDFs, Word, Excel y texto plano.
"""

import json
import csv
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger("jarvis.document_processor")


class DocumentProcessor:
    """Procesa diversos formatos de documentos."""

    def __init__(self):
        logger.info("DocumentProcessor inicializado.")

    # ─── Lectura universal ────────────────────────────────────

    def read_document(self, file_path: str) -> str:
        """Lee un documento y devuelve su contenido como texto."""
        path = Path(file_path)
        if not path.exists():
            return f"Archivo no encontrado: {file_path}"

        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                return self.read_pdf(file_path)
            elif ext in (".docx", ".doc"):
                return self.read_word(file_path)
            elif ext in (".xlsx", ".xls"):
                return self.read_excel(file_path)
            elif ext == ".csv":
                return self.read_csv(file_path)
            elif ext == ".json":
                return self.read_json(file_path)
            elif ext in (".txt", ".md", ".py", ".js", ".html", ".css", ".java",
                          ".cpp", ".c", ".h", ".xml", ".yaml", ".yml", ".ini",
                          ".cfg", ".log", ".bat", ".ps1", ".sh"):
                return self.read_text(file_path)
            else:
                return f"Formato no soportado: {ext}"
        except Exception as e:
            logger.error(f"Error leyendo documento: {e}")
            return f"Error al leer el documento: {e}"

    # ─── PDF ──────────────────────────────────────────────────

    def read_pdf(self, file_path: str) -> str:
        """Lee y extrae texto de un PDF."""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            text_parts = []
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"--- Página {page_num} ---\n{text}")
            doc.close()

            if text_parts:
                full_text = "\n\n".join(text_parts)
                logger.info(f"PDF leído: {len(text_parts)} páginas, {len(full_text)} chars")
                return full_text
            else:
                return "El PDF no contiene texto extraíble (podría ser una imagen escaneada)."

        except ImportError:
            logger.error("PyMuPDF no instalado. pip install PyMuPDF")
            return "Error: PyMuPDF no está instalado."
        except Exception as e:
            logger.error(f"Error leyendo PDF: {e}")
            return f"Error al leer el PDF: {e}"

    def get_pdf_info(self, file_path: str) -> str:
        """Obtiene información metadata de un PDF."""
        try:
            import fitz
            doc = fitz.open(file_path)
            meta = doc.metadata
            info = (
                f"📄 PDF Info: {Path(file_path).name}\n"
                f"  Páginas: {doc.page_count}\n"
                f"  Título: {meta.get('title', 'N/A')}\n"
                f"  Autor: {meta.get('author', 'N/A')}\n"
                f"  Creado: {meta.get('creationDate', 'N/A')}\n"
            )
            doc.close()
            return info
        except Exception as e:
            return f"Error obteniendo info del PDF: {e}"

    # ─── Word (.docx) ────────────────────────────────────────

    def read_word(self, file_path: str) -> str:
        """Lee un documento Word (.docx)."""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Incluir tablas
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(cells))

            text = "\n\n".join(paragraphs)
            logger.info(f"Word leído: {len(paragraphs)} párrafos")
            return text if text else "El documento está vacío."

        except ImportError:
            return "Error: python-docx no está instalado."
        except Exception as e:
            logger.error(f"Error leyendo Word: {e}")
            return f"Error al leer el documento Word: {e}"

    def create_word(self, file_path: str, content: str, title: str = "") -> str:
        """Crea un documento Word con el contenido dado."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt

            doc = Document()

            if title:
                doc.add_heading(title, level=0)

            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            doc.save(file_path)
            return f"Documento Word creado: {file_path}"

        except ImportError:
            return "Error: python-docx no está instalado."
        except Exception as e:
            return f"Error creando documento Word: {e}"

    # ─── Excel (.xlsx) ────────────────────────────────────────

    def read_excel(self, file_path: str) -> str:
        """Lee un archivo Excel y devuelve su contenido como texto."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, read_only=True, data_only=True)
            sheets_text = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows_text = [f"\n📊 Hoja: {sheet_name}"]

                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        rows_text.append(" | ".join(cells))

                sheets_text.append("\n".join(rows_text))

            wb.close()
            text = "\n\n".join(sheets_text)
            return text if text.strip() else "El archivo Excel está vacío."

        except ImportError:
            return "Error: openpyxl no está instalado."
        except Exception as e:
            logger.error(f"Error leyendo Excel: {e}")
            return f"Error al leer el archivo Excel: {e}"

    def create_excel(self, file_path: str, data: list[list], headers: list = None) -> str:
        """Crea un archivo Excel con los datos dados."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font

            wb = Workbook()
            ws = wb.active

            start_row = 1
            if headers:
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = Font(bold=True)
                start_row = 2

            for row_idx, row_data in enumerate(data, start_row):
                for col_idx, value in enumerate(row_data, 1):
                    ws.cell(row=row_idx, column=col_idx, value=value)

            wb.save(file_path)
            return f"Archivo Excel creado: {file_path}"

        except ImportError:
            return "Error: openpyxl no está instalado."
        except Exception as e:
            return f"Error creando Excel: {e}"

    # ─── Texto plano ──────────────────────────────────────────

    def read_text(self, file_path: str) -> str:
        """Lee un archivo de texto plano."""
        try:
            path = Path(file_path)
            text = path.read_text(encoding="utf-8")
            logger.info(f"Archivo texto leído: {len(text)} chars")
            return text if text else "El archivo está vacío."
        except UnicodeDecodeError:
            try:
                text = path.read_text(encoding="latin-1")
                return text
            except Exception as e:
                return f"Error de codificación: {e}"
        except Exception as e:
            return f"Error leyendo archivo de texto: {e}"

    def write_text(self, file_path: str, content: str) -> str:
        """Escribe contenido en un archivo de texto."""
        try:
            path = Path(file_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(content, encoding="utf-8")
            return f"Archivo escrito: {file_path}"
        except Exception as e:
            return f"Error escribiendo archivo: {e}"

    # ─── CSV ──────────────────────────────────────────────────

    def read_csv(self, file_path: str) -> str:
        """Lee un archivo CSV."""
        try:
            rows = []
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(" | ".join(row))
                    if len(rows) > 100:
                        rows.append(f"... (truncado, +{sum(1 for _ in reader)} filas)")
                        break
            return "\n".join(rows) if rows else "El CSV está vacío."
        except Exception as e:
            return f"Error leyendo CSV: {e}"

    # ─── JSON ─────────────────────────────────────────────────

    def read_json(self, file_path: str) -> str:
        """Lee y formatea un archivo JSON."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            return f"Error leyendo JSON: {e}"

    # ─── Resumen de documentos ────────────────────────────────

    def summarize_document(self, file_path: str) -> str:
        """
        Lee un documento y devuelve su contenido para que el LLM lo resuma.
        El resumen real lo hace el orquestador con el LLM.
        """
        content = self.read_document(file_path)
        if content.startswith("Error") or content.startswith("Archivo no"):
            return content

        # Truncar si es muy largo
        max_chars = 10000
        if len(content) > max_chars:
            content = content[:max_chars] + "\n\n[... contenido truncado ...]"

        return content

    def get_document_stats(self, file_path: str) -> str:
        """Obtiene estadísticas de un documento."""
        content = self.read_document(file_path)
        if content.startswith(("Error", "Archivo no")):
            return content

        words = len(content.split())
        chars = len(content)
        lines = content.count("\n") + 1
        paragraphs = content.count("\n\n") + 1

        return (
            f"📊 Estadísticas de {Path(file_path).name}:\n"
            f"  Palabras: {words:,}\n"
            f"  Caracteres: {chars:,}\n"
            f"  Líneas: {lines:,}\n"
            f"  Párrafos: {paragraphs:,}"
        )
